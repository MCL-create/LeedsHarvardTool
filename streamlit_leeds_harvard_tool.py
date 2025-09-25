# streamlit_leeds_harvard_tool.py
"""
Leeds Harvard Referencing Tool — Full production script
Features:
- Extract references from pasted lists, uploaded DOCX/PDF, or webpage URL
- Heuristic parser + Leeds Harvard suggestions
- Narrative feedback report (headings + bullets)
- Exports: report -> DOCX, PDF, XLSX ; reference list -> DOCX, XLSX, TXT, PDF
- Branding header + footer (link to macmillancentreforlearning.co.uk)
- Safe fallbacks for optional libraries
"""

import os
import re
import textwrap
from io import BytesIO
from datetime import datetime
from urllib.parse import urlparse
from pathlib import Path

import streamlit as st
import requests
from bs4 import BeautifulSoup

# Port-binding fix for Render (if provided)
if "PORT" in os.environ:
    try:
        st.set_option("server.port", int(os.environ["PORT"]))
    except Exception:
        pass

# Optional libs with fallbacks
try:
    import docx2txt
except Exception:
    docx2txt = None

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from PyPDF2 import PdfReader
except Exception:
    PdfReader = None

# Required libs assumed present in requirements.txt
from docx import Document as DocxDocument
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
import pandas as pd
import openpyxl

# -------------------------
# Branding colours (user brief)
# -------------------------
BG = "#e6f7f8"
HEADER_BG = "#00a2b3"
HEADER_TEXT = "#ffffff"
MUTED_TEXT = "#37474f"
BORDER = "#80cbc4"
SECONDARY_BG = "#f1f8e9"
LINK = "#0288d1"

# -------------------------
# Page config and CSS
# -------------------------
st.set_page_config(page_title="Leeds Harvard Referencing Tool", page_icon="📚", layout="wide")

st.markdown(
    f"""
    <style>
      :root {{
        --bg: {BG};
        --header-bg: {HEADER_BG};
        --header-text: {HEADER_TEXT};
        --muted-text: {MUTED_TEXT};
        --border: {BORDER};
        --link: {LINK};
      }}
      body {{ background-color: var(--bg); color: var(--muted-text); }}
      .app-header {{ background-color: var(--header-bg); padding: 14px; border-radius: 8px; text-align:center; }}
      .app-footer {{ background-color: {SECONDARY_BG}; padding: 10px; border-radius: 8px; text-align:center; color: var(--muted-text); }}
      .ref-box {{ background: white; border: 1px solid var(--border); padding: 12px; border-radius: 8px; }}
      .muted {{ color: #5f6b6b; font-size:13px; }}
    </style>
    """,
    unsafe_allow_html=True,
)

# -------------------------
# Header (simple, avoids missing logo issues)
# -------------------------
st.markdown(
    f"""
    <div class="app-header">
        <h1 style="color:{HEADER_TEXT}; margin:0;">Leeds Harvard Referencing Checker</h1>
        <div style="color:{HEADER_TEXT}; margin-top:6px; font-size:14px;">Macmillan Centre for Learning</div>
    </div>
    """,
    unsafe_allow_html=True,
)

st.write("")  # spacer

# -------------------------
# Session state
# -------------------------
if "refs" not in st.session_state:
    st.session_state.refs = []         # list of parsed ref dicts

if "report_text" not in st.session_state:
    st.session_state.report_text = ""  # markdown text of last narrative report

if "report_struct" not in st.session_state:
    st.session_state.report_struct = {}  # structured report data

# -------------------------
# Helper: extraction functions
# -------------------------
def safe_extract_docx(uploaded_file):
    """Try python-docx then docx2txt fallback. Returns text or error string."""
    try:
        uploaded_file.seek(0)
        doc = DocxDocument(uploaded_file)  # accepts file-like
        paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n".join(paras)
    except Exception:
        # fallback to docx2txt writing temp file
        if docx2txt:
            try:
                tmp = Path("tmp_uploaded.docx")
                uploaded_file.seek(0)
                tmp.write_bytes(uploaded_file.read())
                txt = docx2txt.process(str(tmp))
                tmp.unlink(missing_ok=True)
                return txt
            except Exception as e:
                return f"[docx extraction error: {e}]"
        return "[Error extracting DOCX: python-docx failed]"

def safe_extract_pdf(uploaded_file):
    """Try PyMuPDF (fitz) then PyPDF2. Returns text or error string."""
    try:
        uploaded_file.seek(0)
        data = uploaded_file.read()
    except Exception:
        return "[Error reading uploaded file]"

    # PyMuPDF first
    if fitz:
        try:
            doc = fitz.open(stream=data, filetype="pdf")
            text_pages = [page.get_text() for page in doc]
            return "\n".join(text_pages)
        except Exception:
            pass
    # PyPDF2 fallback
    if PdfReader:
        try:
            from io import BytesIO
            reader = PdfReader(BytesIO(data))
            pages = []
            for p in reader.pages:
                t = p.extract_text()
                if t:
                    pages.append(t)
            return "\n".join(pages)
        except Exception as e:
            return f"[PDF extraction error: {e}]"
    return "[No PDF extraction library available]"

# -------------------------
# Helper: parsing & checking
# -------------------------
def parse_reference_string(s):
    """
    Heuristic parser: returns dict with keys:
      raw, authors, year, title, source, url, accessed
    """
    raw = (s or "").strip()
    parsed = {"raw": raw, "authors": "", "year": "", "title": "", "source": "", "url": "", "accessed": ""}
    if not raw:
        return parsed

    # URL
    m_url = re.search(r"(https?://\S+)", raw)
    if m_url:
        parsed["url"] = m_url.group(1).rstrip(".,)")
        raw_no_url = raw.replace(m_url.group(1), "").strip()
    else:
        raw_no_url = raw

    # Year (4-digit)
    m_year = re.search(r"\b(19|20)\d{2}\b", raw_no_url)
    if m_year:
        parsed["year"] = m_year.group(0)
        raw_no_year = re.sub(re.escape(m_year.group(0)), "", raw_no_url).strip(" .,:;")
    else:
        raw_no_year = raw_no_url

    # Split by sentence full stops to heuristically assign parts
    parts = [p.strip() for p in re.split(r"\.\s+", raw_no_year) if p.strip()]

    if parts:
        # If first part looks like author (contains comma or 'and'), treat as authors
        if re.search(r"\b[A-Za-z]+,\s*[A-Z]", parts[0]) or re.search(r"\band\b", parts[0], re.I):
            parsed["authors"] = parts[0]
            if len(parts) >= 2:
                parsed["title"] = parts[1]
            if len(parts) >= 3:
                parsed["source"] = ". ".join(parts[2:])
        else:
            # Title-first (e.g., webpages)
            parsed["title"] = parts[0]
            if len(parts) >= 2:
                parsed["source"] = ". ".join(parts[1:])

    # If url exists and no source, use domain
    if parsed["url"] and not parsed["source"]:
        parsed["source"] = urlparse(parsed["url"]).netloc

    # Trim spaces
    for k, v in parsed.items():
        if isinstance(v, str):
            parsed[k] = v.strip()
    return parsed

def check_reference_for_leeds_harvard(parsed):
    """Return list of issues / suggestions (student-facing)."""
    issues = []
    if not parsed.get("authors"):
        issues.append("Add author(s): Leeds Harvard starts with surname then initials (e.g., Smith, J.).")
    if not parsed.get("year"):
        if parsed.get("url"):
            issues.append("No date found — use 'n.d.' for undated web resources and include Accessed date.")
        else:
            issues.append("Add a 4-digit year of publication (e.g., 2023).")
    else:
        if not re.match(r"^(19|20)\d{2}$", parsed["year"]):
            issues.append("Year looks unusual; ensure it is a 4-digit year (e.g., 2023).")
    if not parsed.get("title"):
        issues.append("Add the title of the work. In Leeds Harvard the title is italicised in the reference list.")
    if not parsed.get("source") and not parsed.get("url"):
        issues.append("Add publisher/place (for books) or journal title/volume (for articles).")
    if parsed.get("url") and not parsed.get("accessed"):
        issues.append("Web resource: include an Accessed date (e.g., Accessed: 24 September 2025).")
    return issues

def format_to_leeds(parsed):
    """
    Best-effort Leeds Harvard formatted string.
    For web resources: 'Available at: URL (Accessed: date).'
    """
    authors = parsed.get("authors","").strip()
    year = parsed.get("year","").strip() or "n.d."
    title = parsed.get("title","").strip()
    source = parsed.get("source","").strip()
    url = parsed.get("url","").strip()
    accessed = parsed.get("accessed","").strip()

    if url:
        org = authors or source or urlparse(url).netloc
        title_text = f"*{title}*" if title else ""
        s = f"{org} ({year}) {title_text}."
        s += f" Available at: {url}"
        if accessed:
            s += f" (Accessed: {accessed})."
        else:
            s += " (Accessed: [add date])."
        return s

    # Journal-like heuristics
    if source and ("journal" in source.lower() or "pp." in source.lower() or "," in source):
        header = f"{authors} ({year})" if authors else f"({year})"
        title_part = f" '{title}'." if title else ""
        return f"{header}{title_part} {source}".strip()

    # Book/generic
    header = f"{authors} ({year})." if authors else f"({year})."
    title_part = f" {title}." if title else ""
    source_part = f" {source}." if source else ""
    return f"{header}{title_part}{source_part}".strip()

# -------------------------
# Citation scanning
# -------------------------
def scan_document_for_citations_and_mismatch(text, references):
    """
    Detect simple in-text citation patterns and compare to reference surnames.
    Returns dict: found_citations, referenced_surnames, missing_in_refs, not_cited_in_text.
    """
    if not text:
        return {}
    found = set()
    # (Surname, 2020) or (Surname, 2020a)
    for m in re.finditer(r"\(([^(),\d]+?),\s*(?:[A-Z]?)(19|20)\d{2}\w*\)", text):
        surname = m.group(1).strip().split()[-1]
        found.add(surname.lower())
    # Surname (2020)
    for m in re.finditer(r"\b([A-Z][a-zA-Z'-]+)\s*\((19|20)\d{2}\w*\)", text):
        surname = m.group(1).strip()
        found.add(surname.lower())
    # Build reference surname set
    ref_surnames = set()
    for r in references:
        auth = r.get("authors","")
        if auth:
            if "," in auth:
                surname = auth.split(",")[0].strip()
            else:
                surname = auth.split()[0]
            ref_surnames.add(surname.lower())
    missing_in_refs = sorted(list(found - ref_surnames))
    not_cited = sorted(list(ref_surnames - found))
    return {
        "found_citations": sorted(list(found)),
        "referenced_surnames": sorted(list(ref_surnames)),
        "missing_in_refs": missing_in_refs,
        "not_cited_in_text": not_cited
    }

# -------------------------
# Build narrative report (markdown) + structured
# -------------------------
def build_narrative_report(parsed_refs, citation_results=None):
    lines = []
    struct = {"generated_at": datetime.utcnow().isoformat(), "items": [], "citation_check": citation_results or {}}

    lines.append(f"# Leeds Harvard Referencing Report\n**Generated:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}\n")
    lines.append("## Summary")
    if not parsed_refs:
        lines.append("- No references parsed.")
    else:
        lines.append(f"- {len(parsed_refs)} reference(s) parsed and checked.")
    lines.append("")

    lines.append("## References & Recommendations")
    if not parsed_refs:
        lines.append("No parsed references to display.")
    else:
        for i, pr in enumerate(parsed_refs, start=1):
            issues = check_reference_for_leeds_harvard(pr)
            suggested = format_to_leeds(pr)
            lines.append(f"### {i}. Original")
            lines.append(f"- {pr.get('raw')}")
            lines.append("")
            if issues:
                lines.append("**Issues identified:**")
                for it in issues:
                    lines.append(f"- {it}")
            else:
                lines.append("- No structural issues found (heuristic).")
            lines.append("")
            lines.append("**Suggested Leeds Harvard format:**")
            lines.append(f"- {suggested}")
            lines.append("")
            struct["items"].append({"original": pr.get("raw"), "parsed": pr, "issues": issues, "suggested": suggested})

    lines.append("## Citation vs Reference Check")
    if citation_results:
        found = citation_results.get("found_citations", [])
        missing = citation_results.get("missing_in_refs", [])
        not_cited = citation_results.get("not_cited_in_text", [])
        lines.append(f"- In-text surnames detected (sample): {', '.join(found) if found else 'None detected'}")
        if missing:
            lines.append("**Citations in text not present in the reference list:**")
            for m in missing:
                lines.append(f"- {m} — add a full Leeds Harvard reference.")
        else:
            lines.append("- No missing citations detected (based on heuristics).")
        if not_cited:
            lines.append("**References present in the list NOT cited in text:**")
            for nc in not_cited:
                lines.append(f"- {nc} — check whether this is used or remove.")
    else:
        lines.append("- No citation vs reference check was run.")

    lines.append("")
    lines.append("## Action checklist (student should do):")
    lines.append("- Apply the suggested Leeds Harvard format for each item above.")
    lines.append("- Add 'n.d.' for undated and Accessed dates for web resources.")
    lines.append("- Re-run the check after editing the document to confirm fixes.")
    report_md = "\n\n".join(lines)
    return report_md, struct

# -------------------------
# Export: report -> DOCX, PDF, XLSX
# -------------------------
def export_report_docx(report_md):
    doc = DocxDocument()
    for block in report_md.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            doc.add_heading(block.lstrip("# ").strip(), level=1)
        elif block.startswith("## "):
            doc.add_heading(block.lstrip("# ").strip(), level=2)
        elif block.startswith("### "):
            doc.add_heading(block.lstrip("# ").strip(), level=3)
        else:
            # bullets
            if block.startswith("- "):
                for ln in block.splitlines():
                    if ln.startswith("- "):
                        doc.add_paragraph(ln.lstrip("- ").strip(), style="List Bullet")
                    else:
                        doc.add_paragraph(ln)
            else:
                doc.add_paragraph(block)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def export_report_pdf(report_md):
    bio = BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4, rightMargin=36,leftMargin=36, topMargin=36,bottomMargin=36)
    styles = getSampleStyleSheet()
    flow = []
    for block in report_md.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            flow.append(Paragraph(block.lstrip("# ").strip(), styles["Heading1"]))
        elif block.startswith("## "):
            flow.append(Paragraph(block.lstrip("# ").strip(), styles["Heading2"]))
        elif block.startswith("### "):
            flow.append(Paragraph(block.lstrip("# ").strip(), styles["Heading3"]))
        else:
            for ln in block.splitlines():
                if ln.startswith("- "):
                    flow.append(Paragraph(ln.lstrip("- ").strip(), styles["Bullet"]))
                else:
                    flow.append(Paragraph(ln, styles["Normal"]))
        flow.append(Spacer(1,6))
    doc.build(flow)
    bio.seek(0)
    return bio

def export_report_xlsx(struct):
    bio = BytesIO()
    rows = []
    for item in struct.get("items", []):
        rows.append({
            "Original Reference": item.get("original",""),
            "Issues": "; ".join(item.get("issues", [])),
            "Suggested Leeds Harvard": item.get("suggested","")
        })
    df = pd.DataFrame(rows)
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Report")
    bio.seek(0)
    return bio

# -------------------------
# Export: reference list -> DOCX, XLSX, TXT, PDF
# -------------------------
def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    Add a hyperlink to a python-docx paragraph.
    Returns r_id (not used).
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return r_id

def export_reference_list_docx(parsed_refs):
    doc = DocxDocument()
    doc.add_heading("Reference List (Leeds Harvard)", level=1)
    # sort by surname
    sorted_refs = sorted(parsed_refs, key=lambda r: (r.get("authors","").split(",")[0] if r.get("authors") else ""))
    for pr in sorted_refs:
        p = doc.add_paragraph()
        authors = pr.get("authors","")
        year = pr.get("year","")
        title = pr.get("title","")
        source = pr.get("source","")
        url = pr.get("url","")
        # authors + year
        if authors:
            p.add_run(f"{authors} ")
        if year:
            p.add_run(f"({year}). ")
        # title italic
        if title:
            run = p.add_run(title + ". ")
            run.italic = True
        if source:
            p.add_run(source + ". ")
        if url:
            p.add_run("Available at: ")
            add_hyperlink(p, url, url)
            if pr.get("accessed"):
                p.add_run(f" (Accessed: {pr.get('accessed')}).")
            else:
                p.add_run(" (Accessed: [add date]).")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def export_reference_list_xlsx(parsed_refs):
    bio = BytesIO()
    rows = []
    for pr in parsed_refs:
        rows.append({
            "Authors": pr.get("authors",""),
            "Year": pr.get("year",""),
            "Title": pr.get("title",""),
            "Source": pr.get("source",""),
            "URL": pr.get("url",""),
            "Suggested": format_to_leeds(pr)
        })
    df = pd.DataFrame(rows)
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="References")
    bio.seek(0)
    return bio

def export_reference_list_txt(parsed_refs):
    lines = [pr.get("raw","") for pr in parsed_refs]
    return ("\n".join(lines)).encode("utf-8")

def export_reference_list_pdf(parsed_refs):
    # Simple PDF listing
    bio = BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4)
    styles = getSampleStyleSheet()
    flow = [Paragraph("Reference List (Leeds Harvard)", styles["Heading1"])]
    for pr in parsed_refs:
        text = pr.get("raw","")
        flow.append(Paragraph(text, styles["Normal"]))
    doc.build(flow)
    bio.seek(0)
    return bio

# -------------------------
# UI: Sidebar navigation
# -------------------------
st.sidebar.title("Navigation")
section = st.sidebar.radio("Go to:", ["Check References", "Manual Input", "User Guide"])

# -------------------------
# Page: Check References
# -------------------------
if section == "Check References":
    st.header("Check References — generate narrative report and export")
    st.info("Upload/paste references or a document, run checks, review the narrative report (headings + bullets), then export the exact report to DOCX/PDF/XLSX.")

    mode = st.selectbox("Input method", ["Paste reference list", "Paste full document text", "Upload document (DOCX/PDF)", "Fetch webpage (URL)"])

    # Local parsed refs for this run
    parsed_refs_run = []
    full_text_for_citation = ""

    if mode == "Paste reference list":
        pasted = st.text_area("Paste your reference list here (one reference per line):", height=200)
        if st.button("Parse pasted references"):
            lines = [l.strip() for l in pasted.splitlines() if l.strip()]
            for ln in lines:
                parsed = parse_reference_string(ln)
                parsed_refs_run.append(parsed)
            st.success(f"Parsed {len(parsed_refs_run)} references.")
            for i, p in enumerate(parsed_refs_run, 1):
                st.markdown(f"**{i}.** {p.get('raw')}")
                st.markdown(f"- Suggested: {format_to_leeds(p)}")
                issues = check_reference_for_leeds_harvard(p)
                for it in issues:
                    st.info(it)

    elif mode == "Paste full document text":
        txt = st.text_area("Paste the full document text (or an excerpt):", height=250)
        if st.button("Attempt to extract a reference list from pasted text"):
            # naive heuristic: find 'References' heading
            lines = [ln.strip() for ln in txt.splitlines() if ln.strip()]
            idx = None
            for i, ln in enumerate(lines):
                if re.match(r"^(references|reference list|bibliography)\b", ln, re.I):
                    idx = i+1
                    break
            ref_lines = []
            if idx:
                for ln in lines[idx:]:
                    if len(ln) < 60 and ln.isupper() and not re.search(r"\d", ln):
                        break
                    ref_lines.append(ln)
            if ref_lines:
                st.markdown("Detected reference lines (preview):")
                for ln in ref_lines[:40]:
                    st.write(ln)
                if st.button("Parse detected lines"):
                    for ln in ref_lines:
                        parsed = parse_reference_string(ln)
                        parsed_refs_run.append(parsed)
                    st.success(f"Parsed {len(parsed_refs_run)} references from pasted text.")
            else:
                st.warning("No clear 'References' heading detected. Consider pasting just the reference list or uploading the doc.")

    elif mode == "Upload document (DOCX/PDF)":
        uploaded_file = st.file_uploader("Upload .docx or .pdf", type=["docx", "pdf"])
        if uploaded_file:
            st.info("Extracting text — this may take a moment.")
            if uploaded_file.name.lower().endswith(".pdf"):
                full_text_for_citation = safe_extract_pdf(uploaded_file)
            else:
                full_text_for_citation = safe_extract_docx(uploaded_file)
            st.markdown("**Preview (first 900 chars)**")
            st.text(textwrap.shorten(full_text_for_citation, width=900, placeholder="..."))
            # try to detect references section
            lines = [ln.strip() for ln in full_text_for_citation.splitlines() if ln.strip()]
            idx = None
            for i, ln in enumerate(lines):
                if re.match(r"^(references|reference list|bibliography)\b", ln, re.I):
                    idx = i+1
                    break
            ref_lines = []
            if idx:
                for ln in lines[idx:]:
                    if len(ln) < 60 and ln.isupper() and not re.search(r"\d", ln):
                        break
                    ref_lines.append(ln)
            if ref_lines:
                st.markdown("**Detected Reference section (preview)**")
                for ln in ref_lines[:40]:
                    st.write(ln)
                if st.button("Parse detected reference lines"):
                    for ln in ref_lines:
                        parsed = parse_reference_string(ln)
                        parsed_refs_run.append(parsed)
                    st.success(f"Parsed {len(parsed_refs_run)} references from detected section.")
            else:
                st.warning("No clear 'References' heading detected. You can paste the reference list manually or parse by hand.")

    else:  # Fetch webpage
        url_in = st.text_input("Enter webpage URL (https://...)")
        accessed_in = st.text_input("Accessed date (optional free text, e.g., 24 September 2025)")
        if st.button("Fetch & suggest reference"):
            if not url_in.strip():
                st.warning("Enter a URL.")
            else:
                try:
                    r = requests.get(url_in, timeout=8)
                    s = BeautifulSoup(r.text, "html.parser")
                    title = s.title.string.strip() if s.title and s.title.string else ""
                    site = urlparse(url_in).netloc
                    raw = f"{site} (n.d.) {title}. {url_in}"
                    parsed = {"raw": raw, "authors": site, "year": "n.d.", "title": title, "source": site, "url": url_in, "accessed": accessed_in}
                    parsed_refs_run.append(parsed)
                    st.success("Webpage turned into a suggested reference (review below).")
                    st.markdown(f"- Suggested: {format_to_leeds(parsed)}")
                    for it in check_reference_for_leeds_harvard(parsed):
                        st.info(it)
                except Exception as e:
                    st.error(f"Error fetching URL: {e}")

    # show parsed refs for this run
    if parsed_refs_run:
        st.markdown("### Parsed references (this run)")
        for i, p in enumerate(parsed_refs_run, 1):
            st.markdown(f"**{i}.** {p.get('raw')}")
            st.markdown(f"- Suggested: {format_to_leeds(p)}")
            issues = check_reference_for_leeds_harvard(p)
            for it in issues:
                st.info(it)

    # Citation check (if we have extracted full text)
    citation_results = {}
    if full_text_for_citation:
        if st.button("Run citation vs reference-list check (this run parsed refs)"):
            # prefer parsed_refs_run if available else session refs
            refs_for_check = parsed_refs_run if parsed_refs_run else st.session_state.refs
            citation_results = scan_document_for_citations_and_mismatch(full_text_for_citation, refs_for_check)
            st.json(citation_results)

    # Generate narrative report
    if st.button("Generate narrative report (show full recommendations)"):
        # choose source refs: parsed run OR existing session refs
        source_refs = parsed_refs_run if parsed_refs_run else st.session_state.refs
        if not source_refs:
            st.warning("No references available to build a report. Add references via paste/upload/manual input first.")
        else:
            # if text exists and no citation_results, attempt automatic check
            if full_text_for_citation and not citation_results:
                citation_results = scan_document_for_citations_and_mismatch(full_text_for_citation, source_refs)
            report_md, report_struct = build_narrative_report(source_refs, citation_results)
            st.session_state.report_text = report_md
            st.session_state.report_struct = report_struct
            # persist the refs into session (for exports/reference list)
            st.session_state.refs = source_refs
            st.success("Narrative report generated. Review below.")

    # Show narrative report if present
    if st.session_state.report_text:
        st.markdown("---")
        st.header("Narrative Report (review before export)")
        st.markdown(st.session_state.report_text)
    else:
        st.info("No narrative report yet. Use 'Generate narrative report' once references are available.")

# -------------------------
# Page: Manual Input
# -------------------------
elif section == "Manual Input":
    st.header("Manual reference entry")
    st.info("Add references manually; they are stored in session and used when generating reports.")

    typ = st.selectbox("Type of reference", ["Book", "Journal article", "Report / Grey lit", "Other"])
    if typ == "Book":
        a = st.text_input("Author(s) (Surname, Initials)")
        y = st.text_input("Year")
        t = st.text_input("Title")
        p = st.text_input("Place & Publisher")
        url = st.text_input("URL (optional)")
        accessed = st.text_input("Accessed (optional)")
        if st.button("Save Book reference"):
            raw = f"{a} {y}. {t}. {p}"
            if url:
                raw += f" {url}"
            parsed = {"raw": raw, "authors": a, "year": y, "title": t, "source": p, "url": url, "accessed": accessed}
            st.session_state.refs.append(parsed)
            st.success("Book reference saved to session.")

    elif typ == "Journal article":
        a = st.text_input("Author(s)")
        y = st.text_input("Year")
        t = st.text_input("Article title")
        j = st.text_input("Journal, vol(issue), pp.")
        url = st.text_input("DOI / URL (optional)")
        if st.button("Save Journal reference"):
            raw = f"{a} {y}. {t}. {j}"
            if url:
                raw += f" {url}"
            parsed = {"raw": raw, "authors": a, "year": y, "title": t, "source": j, "url": url}
            st.session_state.refs.append(parsed)
            st.success("Journal reference saved to session.")

    elif typ == "Report / Grey lit":
        a = st.text_input("Organisation / Author")
        y = st.text_input("Year")
        t = st.text_input("Title")
        src = st.text_input("Publisher / Source")
        url = st.text_input("URL (optional)")
        accessed = st.text_input("Accessed (optional)")
        if st.button("Save Report"):
            raw = f"{a} {y}. {t}. {src}"
            if url:
                raw += f" {url}"
            parsed = {"raw": raw, "authors": a, "year": y, "title": t, "source": src, "url": url, "accessed": accessed}
            st.session_state.refs.append(parsed)
            st.success("Report saved.")

    else:
        raw = st.text_area("Full reference (paste free text)")
        if st.button("Save Other"):
            parsed = parse_reference_string(raw)
            st.session_state.refs.append(parsed)
            st.success("Reference saved.")

    st.markdown("---")
    st.subheader("Current saved references (session)")
    if st.session_state.refs:
        for i, r in enumerate(st.session_state.refs, 1):
            st.markdown(f"**{i}.** {r.get('raw')}")
            st.markdown(f"- Suggested: {format_to_leeds(r)}")
    else:
        st.info("No references saved in this session. Add some above or use Check References.")

# -------------------------
# Page: User Guide
# -------------------------
elif section == "User Guide":
    st.header("User Guide — how to use this tool")
    st.markdown("""
    **What this tool does**
    - Parses reference lists (or detects references in uploaded documents), identifies structural issues according to Leeds Harvard style, and produces a narrative feedback report with suggested corrections.

    **Recommended workflow for students**
    1. Use *Check References* → upload your document or paste the reference list.
    2. Click **Generate narrative report**.
    3. Read the report, implement corrections, re-run if needed.
    4. Export the final narrative report to DOCX/PDF/XLSX and include it as feedback or evidence.

    **For tutors**
    - Use the generated narrative report as a feedback sheet. The tool intentionally does not automatically rewrite students' work — students learn by editing themselves.
    """)
    st.markdown("---")
    st.markdown(f"<div class='muted'>For exact Leeds Harvard local guidance, see: <a href='https://library.leeds.ac.uk/info/1404/referencing/46/harvard_style' target='_blank'>Leeds Harvard guidance</a></div>", unsafe_allow_html=True)

# -------------------------
# Always-visible Export area (bottom)
# -------------------------
st.markdown("---")
st.subheader("Export area — download the final narrative report and reference list")

col1, col2 = st.columns(2)

with col1:
    st.markdown("### Narrative report")
    if st.session_state.report_text:
        docx_buf = export_report_docx(st.session_state.report_text)
        pdf_buf = export_report_pdf(st.session_state.report_text)
        xlsx_buf = export_report_xlsx(st.session_state.report_struct)
        st.download_button("Download report (.docx)", data=docx_buf.getvalue(), file_name="LeedsHarvard_Report.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.download_button("Download report (.pdf)", data=pdf_buf.getvalue(), file_name="LeedsHarvard_Report.pdf",
                           mime="application/pdf")
        st.download_button("Download report (.xlsx)", data=xlsx_buf.getvalue(), file_name="LeedsHarvard_Report.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    else:
        st.info("No narrative report available. Generate one from Check References.")

with col2:
    st.markdown("### Reference list exports")
    if st.session_state.refs:
        # reference list docx
        rl_docx = export_reference_list_docx(st.session_state.refs)
        rl_xlsx = export_reference_list_xlsx(st.session_state.refs)
        rl_txt = export_reference_list_txt(st.session_state.refs)
        rl_pdf = export_reference_list_pdf(st.session_state.refs)
        st.download_button("Download reference list (.docx)", data=rl_docx.getvalue(), file_name="Reference_List.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.download_button("Download reference list (.xlsx)", data=rl_xlsx.getvalue(), file_name="Reference_List.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        st.download_button("Download reference list (.txt)", data=rl_txt, file_name="Reference_List.txt",
                           mime="text/plain")
        st.download_button("Download reference list (.pdf)", data=rl_pdf.getvalue(), file_name="Reference_List.pdf",
                           mime="application/pdf")
    else:
        st.info("No references saved in session. Add references to enable exports.")

# -------------------------
# Footer
# -------------------------
st.markdown(
    f"""
    <div class="app-footer">
      <div>© {datetime.now().year} <a href="https://macmillancentreforlearning.co.uk" target="_blank" style="color:{LINK}; text-decoration:none;">Macmillan Centre for Learning</a> | Leeds Harvard Referencing Tool</div>
    </div>
    """,
    unsafe_allow_html=True,
)
