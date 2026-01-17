import streamlit as st
import re
import os
from io import BytesIO
from docx import Document
import leeds_harvard_tool as lht

if 'bibliography' not in st.session_state: 
    st.session_state.bibliography = []

st.set_page_config(page_title="MCL Referencing Assistant", page_icon="🎓", layout="wide")

st.markdown("""
    <style>
    .stApp { background-color: #e6f7f8; } 
    div.stButton > button { background-color: #009688; color: white; border-radius: 5px; font-weight: bold; width: 100%; }
    .stTabs [aria-selected="true"] { background-color: #009688 !important; color: white !important; }
    .content-box { background-color: white; padding: 25px; border-radius: 10px; border-left: 5px solid #009688; margin-bottom: 20px; }
    .glossary-term { color: #009688; font-weight: bold; font-size: 1.2em; margin-top: 20px; }
    .example-box { background-color: #f9f9f9; padding: 15px; border: 1px dashed #009688; border-radius: 5px; }
    </style>
""", unsafe_allow_html=True)
if os.path.exists("assets/Header.png"): 
    st.image("assets/Header.png", use_column_width=True)

tabs = st.tabs(["🏠 Guide", "📖 Book", "📰 Journal", "🌐 Website", "📋 Bibliography", "🔍 Smart Audit", "📚 Glossary"])

# --- TAB 1: FULL GUIDE (COMPLETE TEXT FROM SCREENSHOTS) ---
with tabs[0]:
    st.title("🎓 Leeds Harvard Referencing Guide")
    st.markdown("""
    <div class="content-box">
    <p>Below are clear examples of how three common source types are presented using the Leeds Harvard style. This referencing system prioritises author–date citation, consistency, and sufficient information to enable the reader to retrieve the source. Each example shows (a) how to format the full reference and (b) how the matching in-text citation appears within academic work.</p>
    
    <h4>Book example</h4>
    <p><strong>Full reference</strong><br>
    Smith, J. (2022) <em>Understanding professional practice.</em> 2nd edn. London: Routledge.</p>
    <p><strong>In-text citation</strong><br>
    Professional practice requires reflection and ongoing learning (Smith, 2022).</p>

    <h4>Journal article example</h4>
    <p><strong>Full reference</strong><br>
    Brown, L. and Green, T. (2023) ‘Developing reflective capacity in vocational education’, <em>Journal of Education and Work</em>, 36(4), pp. 415–431. https://doi.org/10.1080/13639080.2023.1234567</p>
    <p><strong>In-text citation</strong><br>
    Reflective capacity is argued to be central to effective professional development (Brown and Green, 2023).</p>

    <h4>Website example</h4>
    <p><strong>Full reference</strong><br>
    Scottish Social Services Council (2024) <em>Codes of practice for social service workers and employers.</em> Available at: https://www.sssc.uk.com/codes-of-practice (Accessed: 13 January 2026).</p>
    <p><strong>In-text citation</strong><br>
    The most recent codes emphasise accountability and high-quality care (Scottish Social Services Council, 2024).</p>

    <hr>
    <h3>Quick Start Instructions</h3>
    <ul>
        <li>Enter source details in the <strong>Book, Journal,</strong> or <strong>Website</strong> tabs.</li>
        <li>Review and alphabetise your list in the <strong>Bibliography</strong> tab.</li>
        <li>Use the <strong>Smart Audit</strong> tool to ensure your essay citations match your list.</li>
    </ul>
    </div>
    """, unsafe_allow_html=True)
    
    # PRINTABLE GUIDE EXPORT (FIXED SYNTAX)
    doc_g = Document()
    doc_g.add_heading('MCL Reference Guide', 0)
    doc_g.add_paragraph('Core Formatting: Family name, INITIAL(S). Year. Title. Place: Publisher.')
    doc_g.add_heading('Examples', level=1)
    doc_g.add_paragraph('Book: Smith, J. (2022) Understanding professional practice. London: Routledge.')
    doc_g.add_paragraph('Journal: Brown, L. and Green, T. (2023) ‘Developing reflective capacity’, Journal of Education and Work, 36(4), pp. 415–431.')
    buf_g = BytesIO(); doc_g.save(buf_g)
    st.download_button("🖨️ Download Full Printable Guide (.docx)", buf_g.getvalue(), "Leeds_Harvard_Full_Guide.docx")

# --- TAB 2: BOOK ---
with tabs[1]:
    st.header("📖 Add a Book")
    with st.form("b_form"):
        ba=st.text_input("Author"); by=st.text_input("Year"); bt=st.text_input("Title"); bp=st.text_input("Publisher")
        if st.form_submit_button("Add Book"):
            st.session_state.bibliography.append(lht.generate_book_reference(ba,by,bt,bp))
            st.success("Book Added")

# --- TAB 3: JOURNAL (RESTORED) ---
with tabs[2]:
    st.header("📰 Add a Journal Article")
    with st.form("j_form"):
        ja=st.text_input("Author"); jy=st.text_input("Year"); jt=st.text_input("Article Title"); jn=st.text_input("Journal Name"); jv=st.text_input("Vol"); ji=st.text_input("Issue"); jp=st.text_input("Pages")
        if st.form_submit_button("Add Journal"):
            st.session_state.bibliography.append(lht.generate_journal_reference(ja,jy,jt,jn,jv,ji,jp))
            st.success("Journal Added")

# --- TAB 4: WEBSITE (RESTORED) ---
with tabs[3]:
    st.header("🌐 Add a Website")
    with st.form("w_form"):
        wa=st.text_input("Author/Org"); wy=st.text_input("Year"); wt=st.text_input("Page Title"); wu=st.text_input("URL"); wd=st.text_input("Access Date")
        if st.form_submit_button("Add Website"):
            st.session_state.bibliography.append(lht.generate_web_reference(wa,wy,wt,wu,wd))
            st.success("Website Added")

# --- TAB 5: BIBLIOGRAPHY ---
with tabs[4]:
    st.header("📋 Your Bibliography")
    if st.session_state.bibliography:
        if st.button("🪄 One-Click Correction"):
            st.session_state.bibliography = lht.apply_one_click_corrections(st.session_state.bibliography)
            st.rerun()
        st.session_state.bibliography.sort()
        for r in st.session_state.bibliography: st.markdown(f"- {r}", unsafe_allow_html=True)
        doc_b = Document(); doc_b.add_heading('Bibliography', 0)
        for r in st.session_state.bibliography:
            p = doc_b.add_paragraph()
            parts = re.split(r'(<i>.*?</i>)', r)
            for part in parts:
                if part.startswith('<i>'): p.add_run(part.replace('<i>','').replace('</i>','')).italic = True
                else: p.add_run(part)
        buf_b = BytesIO(); doc_b.save(buf_b); st.download_button("📥 Download Bibliography (.docx)", buf_b.getvalue(), "MCL_Bibliography.docx")

# --- TAB 6: SMART AUDIT ---
with tabs[5]:
    st.header("🔍 Smart Essay Audit")
    up = st.file_uploader("Upload Essay (.docx)", type="docx")
    if up:
        text = lht.extract_text_from_docx(up)
        if text:
            clean_bib = [lht.clean_text(b) for b in st.session_state.bibliography]
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            results = []
            for i, p in enumerate(paragraphs):
                cites = re.findall(r'\(([^)]{2,100}?\d{4}[^)]{0,50}?)\)', p)
                for c in cites:
                    clean_cite = lht.clean_text(c)
                    matched = any(clean_cite in cb for cb in clean_bib if cb)
                    results.append({"Para": i+1, "Citation": f"({c})", "Status": "✅" if matched else "❌", "Feedback": "Verified" if matched else "⚠️ Missing"})
            if results: st.table(results)
            # Narrative Report Export
            doc_r = Document(); doc_r.add_heading('Audit Report', 0)
            for res in results: doc_r.add_paragraph(f"Para {res['Para']}: {res['Citation']} - {res['Feedback']}")
            buf_r = BytesIO(); doc_r.save(buf_r); st.download_button("📊 Download Audit Report", buf_r.getvalue(), "Audit_Report.docx")

# --- TAB 7: GLOSSARY (RESTORED) ---
with tabs[6]:
    st.header("Glossary of Key Academic Writing Terms")
    st.markdown("""
    <div class="content-box">
        <div class="glossary-term">Plagiarism</div>
        <p><strong>Definition:</strong> Plagiarism is the act of presenting another person’s ideas, words, data, or creative work as one’s own without appropriate acknowledgement (QAA, 2019).</p>
        <p>In the Scottish context, this is a breach of professional values consistent with the <strong>SSSC Codes of Practice (2024)</strong>.</p>
        <div class="example-box">
            <strong>Original:</strong> “Assessment feedback plays a critical role...” (Nicol and Macfarlane-Dick, 2006).<br>
            <strong>Verdict:</strong> Plagiarism if copied exactly with no quotation marks or citation.
        </div>
        <div class="glossary-term">Paraphrasing</div>
        <p>Restating another author's ideas in your own words while accurately preserving the original meaning (Pears and Shields, 2022).</p>
        <div class="glossary-term">Direct Quote</div>
        <p>Using exact words, enclosed in quotation marks, with a citation including a page number (Cottrell, 2019).</p>
    </div>
    """, unsafe_allow_html=True)
